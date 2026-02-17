"""
OCR Processing Module for Resume Text Extraction
Uses Tesseract OCR (pytesseract) for reliable text extraction from PDF and DOCX files
"""

import os
import tempfile
from typing import Optional, Tuple
from pdf2image import convert_from_bytes
from PIL import Image
import pytesseract
import io


class OCRProcessor:
    """Process PDF and DOCX files to extract text using Tesseract OCR"""
    
    def __init__(self):
        """Initialize OCR processor"""
        # Tesseract doesn't need initialization, just needs to be installed on system
        # Check if tesseract is available
        try:
            pytesseract.get_tesseract_version()
            print("[OCR] Tesseract OCR initialized successfully")
        except Exception as e:
            print(f"[OCR] Warning: Tesseract might not be installed: {e}")
    
    def process_file(self, file_data: bytes, file_type: str) -> Tuple[bool, Optional[str], str]:
        """
        Process a file and extract text using OCR
        
        Args:
            file_data: Binary file content
            file_type: File extension ('pdf' or 'docx')
            
        Returns:
            Tuple of (success, extracted_text, error_message)
        """
        try:
            if file_type.lower() == 'pdf':
                return self._process_pdf(file_data)
            elif file_type.lower() == 'docx':
                return self._process_docx(file_data)
            else:
                return False, None, f"Unsupported file type: {file_type}"
        except Exception as e:
            return False, None, f"OCR processing error: {str(e)}"
    
    def _process_pdf(self, pdf_data: bytes) -> Tuple[bool, Optional[str], str]:
        """Convert PDF to images and extract text from each page"""
        try:
            print(f"[OCR] Starting PDF processing, file size: {len(pdf_data)} bytes")
            
            # Convert PDF bytes to images (one image per page)
            images = convert_from_bytes(pdf_data)
            print(f"[OCR] Converted PDF to {len(images)} page(s)")
            
            all_text = []
            for page_num, image in enumerate(images, start=1):
                print(f"[OCR] Processing page {page_num}...")
                
                # Extract text from image using Tesseract
                text = self._extract_text_from_image(image)
                print(f"[OCR] Page {page_num} extracted {len(text)} characters")
                
                if text.strip():  # Only add if there's actual text
                    all_text.append(f"--- Page {page_num} ---\n{text}")
            
            if all_text:
                combined_text = "\n\n".join(all_text)
                print(f"[OCR] ✅ PDF OCR SUCCESS: Total {len(combined_text)} characters extracted")
                return True, combined_text, "PDF OCR completed successfully"
            else:
                print(f"[OCR] ❌ PDF OCR FAILED: No text extracted from any page")
                return False, None, "No text extracted from PDF"
                
        except Exception as e:
            print(f"[OCR] ❌ PDF processing error: {str(e)}")
            import traceback
            traceback.print_exc()
            return False, None, f"PDF processing error: {str(e)}"
    
    def _process_docx(self, docx_data: bytes) -> Tuple[bool, Optional[str], str]:
        """Extract text directly from DOCX (no OCR needed for text)"""
        try:
            from docx import Document
            
            print(f"[OCR] Starting DOCX processing, file size: {len(docx_data)} bytes")
            
            # Save DOCX to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
                tmp_docx.write(docx_data)
                tmp_docx_path = tmp_docx.name
            
            try:
                # Open the DOCX file
                doc = Document(tmp_docx_path)
                
                # Extract text directly from DOCX (faster than OCR)
                docx_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        docx_text.append(para.text)
                
                # Also check tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                docx_text.append(cell.text)
                
                if docx_text:
                    combined_text = "\n".join(docx_text)
                    print(f"[OCR] ✅ DOCX text extracted: {len(combined_text)} characters")
                    return True, combined_text, "DOCX text extracted successfully"
                else:
                    print(f"[OCR] ❌ No text found in DOCX")
                    return False, None, "No text found in DOCX"
                    
            finally:
                # Clean up temporary file
                if os.path.exists(tmp_docx_path):
                    os.unlink(tmp_docx_path)
                    
        except Exception as e:
            print(f"[OCR] ❌ DOCX processing error: {str(e)}")
            import traceback
            traceback.print_exc()
            return False, None, f"DOCX processing error: {str(e)}"
    
    def _extract_text_from_image(self, image: Image.Image) -> str:
        """Extract text from PIL Image using Tesseract OCR"""
        try:
            print(f"[OCR] Extracting text from image, size: {image.size}, mode: {image.mode}")
            
            # Use Tesseract to extract text
            # config parameter optimizes for different scenarios
            text = pytesseract.image_to_string(image, config='--psm 6')
            
            extracted_lines = [line.strip() for line in text.split('\n') if line.strip()]
            print(f"[OCR] Extracted {len(extracted_lines)} non-empty lines")
            
            return text
            
        except Exception as e:
            print(f"[OCR] ❌ Error extracting text from image: {e}")
            import traceback
            traceback.print_exc()
            return ""
