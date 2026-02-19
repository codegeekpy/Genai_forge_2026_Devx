"""
Resume Builder Module — PDF/DOCX generation + AI text optimization
Ported from ai_resume_template- and adapted for TalentForge
"""

from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import httpx
from pydantic import BaseModel
from typing import List, Optional, Dict


# ── Pydantic Models ──

class PersonalInfo(BaseModel):
    fullName: Optional[str] = None
    jobTitle: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    linkedin: Optional[str] = None
    website: Optional[str] = None
    summary: Optional[str] = None


class ExperienceEntry(BaseModel):
    title: str = ""
    company: str = ""
    date: str = ""
    description: str = ""


class EducationEntry(BaseModel):
    degree: str = ""
    school: str = ""
    date: str = ""


class ResumeBuilderData(BaseModel):
    personal: PersonalInfo = PersonalInfo()
    experience: List[ExperienceEntry] = []
    education: List[EducationEntry] = []
    skills: List[str] = []
    template: str = "modern"
    color: str = "#0e6b5e"  # TalentForge teal by default
    font: str = "sans"


class TextOptimizeRequest(BaseModel):
    text: str


# ── Color definitions ──

COLORS = {
    "teal": "#0e6b5e",
    "navy": "#0d2137",
    "gold": "#b8860b",
    "blue": "#2563EB",
    "emerald": "#059669",
    "purple": "#7C3AED",
    "slate": "#475569",
    "red": "#DC2626",
    "black": "#000000",
    "orange": "#EA580C",
}


# ═══════════════════════════════════════════
#  PDF GENERATION (fpdf2)
# ═══════════════════════════════════════════

class ResumePDF(FPDF):
    def __init__(self, layout='modern', color_hex='#0e6b5e', font_type='sans'):
        super().__init__()
        self.layout_type = layout
        self.primary_color = self._hex_to_rgb(color_hex)
        self.font_type = font_type

    def _hex_to_rgb(self, hex_str):
        hex_str = hex_str.lstrip('#')
        return tuple(int(hex_str[i:i+2], 16) for i in (0, 2, 4))

    def setup_font(self, style='', size=11):
        font_map = {
            'sans': 'helvetica',
            'serif': 'times',
            'mono': 'courier'
        }
        self.set_font(font_map.get(self.font_type, 'helvetica'), style=style, size=size)


def generate_pdf(data: dict) -> bytes:
    """Generate a styled PDF resume from builder data."""
    layout = data.get('template', 'modern')
    color_hex = data.get('color', '#0e6b5e')
    font_type = data.get('font', 'sans')

    pdf = ResumePDF(layout=layout, color_hex=color_hex, font_type=font_type)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    personal = data.get('personal', {})

    # ── Header by layout ──
    if layout == 'modern':
        pdf.setup_font("B", 24)
        pdf.set_text_color(*pdf.primary_color)
        pdf.cell(0, 12, personal.get('fullName', 'Name'), ln=True)

        pdf.setup_font("I", 14)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 10, personal.get('jobTitle', 'Title'), ln=True)

        pdf.setup_font("", 10)
        pdf.set_text_color(80, 80, 80)
        parts = [personal.get('email', ''), personal.get('phone', ''), personal.get('linkedin', '')]
        contact = " | ".join(p for p in parts if p)
        pdf.cell(0, 8, contact, ln=True)
        pdf.ln(5)

        pdf.set_draw_color(*pdf.primary_color)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(8)

    elif layout == 'classic':
        pdf.font_type = 'serif'
        pdf.setup_font("B", 20)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 10, personal.get('fullName', 'Name').upper(), align="C", ln=True)

        pdf.setup_font("", 11)
        pdf.cell(0, 6, personal.get('jobTitle', ''), align="C", ln=True)

        pdf.setup_font("", 10)
        parts = [personal.get('email', ''), personal.get('phone', ''), personal.get('linkedin', '')]
        contact = "  |  ".join(p for p in parts if p)
        pdf.cell(0, 6, contact, align="C", ln=True)
        pdf.ln(5)
        pdf.line(40, pdf.get_y(), 170, pdf.get_y())
        pdf.ln(8)

    elif layout == 'executive':
        pdf.setup_font("B", 22)
        pdf.set_text_color(*pdf.primary_color)
        pdf.cell(0, 12, personal.get('fullName', 'Name'), ln=True)

        pdf.setup_font("I", 12)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(0, 8, personal.get('jobTitle', ''), ln=True)

        pdf.setup_font("", 9)
        pdf.set_text_color(100, 100, 100)
        parts = [personal.get('email', ''), personal.get('phone', ''), personal.get('linkedin', '')]
        contact = " | ".join(p for p in parts if p)
        pdf.cell(0, 7, contact, ln=True)
        pdf.ln(4)

        pdf.set_draw_color(*pdf.primary_color)
        pdf.set_line_width(0.8)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.set_line_width(0.2)
        pdf.ln(8)

    elif layout == 'creative':
        pdf.setup_font("B", 28)
        pdf.set_text_color(*pdf.primary_color)
        name_parts = personal.get('fullName', 'Name').split(' ', 1)
        pdf.cell(0, 14, name_parts[0] if name_parts else 'Name', ln=True)
        if len(name_parts) > 1:
            pdf.setup_font("B", 28)
            pdf.set_text_color(30, 30, 30)
            pdf.cell(0, 14, name_parts[1], ln=True)

        pdf.setup_font("B", 12)
        pdf.set_text_color(255, 255, 255)
        pdf.set_fill_color(30, 30, 30)
        pdf.cell(0, 10, f"  {personal.get('jobTitle', '').upper()}  ", ln=True, fill=True)
        pdf.ln(5)

        pdf.setup_font("", 9)
        pdf.set_text_color(100, 100, 100)
        parts = [personal.get('email', ''), personal.get('phone', ''), personal.get('linkedin', '')]
        contact = " | ".join(p for p in parts if p)
        pdf.cell(0, 7, contact, ln=True)
        pdf.ln(6)

    else:  # minimalist
        pdf.setup_font("B", 30)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 15, personal.get('fullName', 'Name'), ln=True)
        pdf.setup_font("", 12)
        pdf.set_text_color(*pdf.primary_color)
        pdf.cell(0, 10, personal.get('jobTitle', '').upper(), ln=True)

        pdf.setup_font("", 9)
        pdf.set_text_color(130, 130, 130)
        parts = [personal.get('email', ''), personal.get('phone', '')]
        contact = " | ".join(p for p in parts if p)
        pdf.cell(0, 7, contact, ln=True)
        pdf.ln(10)

    # ── Summary ──
    if personal.get('summary'):
        _section_heading(pdf, "SUMMARY", layout)
        pdf.setup_font("", 10)
        pdf.set_text_color(50, 50, 50)
        pdf.multi_cell(0, 6, personal['summary'])
        pdf.ln(6)

    # ── Experience ──
    if data.get('experience'):
        _section_heading(pdf, "EXPERIENCE", layout)
        for exp in data['experience']:
            pdf.setup_font("B", 11)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(140, 7, exp.get('title', ''))
            pdf.setup_font("I", 10)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(0, 7, exp.get('date', ''), align="R", ln=True)

            pdf.setup_font("B", 10)
            pdf.set_text_color(70, 70, 70)
            pdf.cell(0, 6, exp.get('company', ''), ln=True)

            pdf.setup_font("", 10)
            pdf.set_text_color(50, 50, 50)
            pdf.multi_cell(0, 5, exp.get('description', ''))
            pdf.ln(4)
        pdf.ln(4)

    # ── Education ──
    if data.get('education'):
        _section_heading(pdf, "EDUCATION", layout)
        for edu in data['education']:
            pdf.setup_font("B", 10)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(140, 6, edu.get('degree', ''))
            pdf.setup_font("", 10)
            pdf.cell(0, 6, edu.get('date', ''), align="R", ln=True)
            pdf.setup_font("", 10)
            pdf.set_text_color(70, 70, 70)
            pdf.cell(0, 5, edu.get('school', ''), ln=True)
            pdf.ln(2)
        pdf.ln(4)

    # ── Skills ──
    if data.get('skills'):
        _section_heading(pdf, "SKILLS", layout)
        pdf.setup_font("", 10)
        pdf.set_text_color(50, 50, 50)
        skills_text = ", ".join(s for s in data['skills'] if s)
        pdf.multi_cell(0, 6, skills_text)

    return bytes(pdf.output())


def _section_heading(pdf: ResumePDF, title: str, layout: str):
    """Render a section heading styled per layout."""
    pdf.setup_font("B", 12)
    pdf.set_text_color(*pdf.primary_color)

    if layout == 'classic':
        pdf.cell(0, 8, title, ln=True)
        pdf.set_draw_color(*pdf.primary_color)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(4)
    elif layout == 'creative':
        pdf.set_draw_color(*pdf.primary_color)
        pdf.set_line_width(0.6)
        y = pdf.get_y() + 4
        pdf.line(10, y, 30, y)
        pdf.set_line_width(0.2)
        pdf.cell(25)
        pdf.cell(0, 8, title, ln=True)
    else:
        pdf.cell(0, 8, title, ln=True)

    pdf.ln(2)


# ═══════════════════════════════════════════
#  DOCX GENERATION (python-docx)
# ═══════════════════════════════════════════

def generate_docx(data: dict) -> io.BytesIO:
    """Generate a styled DOCX resume from builder data."""
    layout = data.get('template', 'modern')
    color_hex = data.get('color', '#0e6b5e')
    document = Document()

    personal = data.get('personal', {})

    # Font defaults
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman' if layout == 'classic' else 'Arial'
    font.size = Pt(11)

    # Parse color
    c = color_hex.lstrip('#')
    r, g, b = int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16)
    accent_color = RGBColor(r, g, b)

    fullName = personal.get('fullName', 'Your Name')

    # ── Header ──
    if layout == 'classic':
        h1 = document.add_paragraph()
        run = h1.add_run(fullName.upper())
        run.bold = True
        run.font.size = Pt(22)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if personal.get('jobTitle'):
            run_title = p.add_run(f"{personal['jobTitle']}\n")
            run_title.italic = True
        parts = [personal.get('email', ''), personal.get('phone', ''), personal.get('linkedin', '')]
        p.add_run(" | ".join(p2 for p2 in parts if p2))
    else:
        h1 = document.add_heading(fullName, 0)
        for run in h1.runs:
            run.font.color.rgb = accent_color

        p = document.add_paragraph()
        if personal.get('jobTitle'):
            p.add_run(f"{personal['jobTitle']}\n").bold = True
        parts = [personal.get('email', ''), personal.get('phone', '')]
        p.add_run(" | ".join(p2 for p2 in parts if p2))

    # ── Section helper ──
    def add_section_header(text):
        h = document.add_heading(text, level=1)
        for run in h.runs:
            run.font.color.rgb = accent_color

    # ── Summary ──
    if personal.get('summary'):
        add_section_header('Profile Summary')
        document.add_paragraph(personal['summary'])

    # ── Experience ──
    if data.get('experience'):
        add_section_header('Professional Experience')
        for exp in data['experience']:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            title_run = p.add_run(exp.get('title', ''))
            title_run.bold = True
            if exp.get('date'):
                p.add_run(f"\t{exp['date']}").italic = True

            if exp.get('company'):
                p2 = document.add_paragraph()
                p2.add_run(exp['company']).italic = True

            if exp.get('description'):
                document.add_paragraph(exp['description'])

    # ── Education ──
    if data.get('education'):
        add_section_header('Education')
        for edu in data['education']:
            p = document.add_paragraph()
            p.add_run(edu.get('degree', '')).bold = True
            school = edu.get('school', '')
            date = edu.get('date', '')
            if school or date:
                p.add_run(f", {school} ({date})")

    # ── Skills ──
    if data.get('skills'):
        add_section_header('Skills')
        skills_text = ', '.join(s for s in data['skills'] if s)
        document.add_paragraph(skills_text)

    docx_stream = io.BytesIO()
    document.save(docx_stream)
    docx_stream.seek(0)
    return docx_stream


# ═══════════════════════════════════════════
#  AI TEXT OPTIMIZATION (Groq)
# ═══════════════════════════════════════════

def ai_optimize_text(text: str) -> str:
    """Polish and improve text using Groq LLM."""
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return text  # Return original if no key

    prompt = (
        "Please polish and improve the following text to be more professional "
        "and impactful for a resume. Return ONLY the improved text, no explanations:\n\n"
        f"{text}"
    )

    try:
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "llama-3.3-70b-versatile",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3,
            "max_tokens": 500
        }
        with httpx.Client(timeout=30.0) as client:
            response = client.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers=headers, json=payload
            )
            if response.status_code == 200:
                result = response.json()
                return result["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print(f"[ResumeBuilder] AI optimize error: {e}")

    return text


def ai_generate_summary(data: dict) -> str:
    """Generate a professional resume summary from resume data."""
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return ""

    personal = data.get('personal', {})
    context = f"Name: {personal.get('fullName', 'Unknown')}, Job Title: {personal.get('jobTitle', 'Professional')}. "
    experience = data.get('experience', [])
    if experience:
        context += "Experience: " + ", ".join(
            f"{e.get('title', '')} at {e.get('company', '')}" for e in experience
        )
    skills = data.get('skills', [])
    if skills:
        context += f". Skills: {', '.join(skills[:10])}"

    prompt = (
        f"Write a professional resume summary (3-4 sentences) for this candidate: {context}. "
        "Return ONLY the summary text, no explanations or formatting."
    )

    try:
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "llama-3.3-70b-versatile",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.4,
            "max_tokens": 300
        }
        with httpx.Client(timeout=30.0) as client:
            response = client.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers=headers, json=payload
            )
            if response.status_code == 200:
                result = response.json()
                return result["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print(f"[ResumeBuilder] AI summary error: {e}")

    return ""


# ═══════════════════════════════════════════
#  PRE-FILL FROM EXTRACTED RESUME DATA
# ═══════════════════════════════════════════

def transform_extracted_to_builder(extracted_info: dict, user_name: str = "") -> dict:
    """
    Convert extracted resume JSON (from Groq extraction) to resume builder format.

    Extracted format (from groq_extractor):
        {name, email, phone, education[{degree,institution,year,details}],
         skills[], experience[{title,company,duration,responsibilities}],
         projects[{name,description,technologies,duration}], certifications[]}

    Builder format:
        {personal:{fullName,jobTitle,email,phone,linkedin,summary},
         experience[{title,company,date,description}],
         education[{degree,school,date}], skills[]}
    """
    if not extracted_info:
        return {
            "personal": {"fullName": user_name},
            "experience": [], "education": [], "skills": []
        }

    # Personal info
    personal = {
        "fullName": extracted_info.get("name", "") or user_name,
        "jobTitle": "",
        "email": extracted_info.get("email", ""),
        "phone": extracted_info.get("phone", ""),
        "linkedin": "",
        "website": "",
        "summary": ""
    }

    # Try to infer job title from latest experience
    exp_list = extracted_info.get("experience", [])
    if exp_list and isinstance(exp_list, list) and len(exp_list) > 0:
        personal["jobTitle"] = exp_list[0].get("title", "")

    # Experience
    experience = []
    for exp in exp_list:
        if isinstance(exp, dict):
            experience.append({
                "title": exp.get("title", ""),
                "company": exp.get("company", ""),
                "date": exp.get("duration", ""),
                "description": exp.get("responsibilities", "")
            })

    # Education
    education = []
    for edu in extracted_info.get("education", []):
        if isinstance(edu, dict):
            education.append({
                "degree": edu.get("degree", ""),
                "school": edu.get("institution", ""),
                "date": edu.get("year", "")
            })

    # Skills
    skills = extracted_info.get("skills", [])
    if not isinstance(skills, list):
        skills = []

    return {
        "personal": personal,
        "experience": experience,
        "education": education,
        "skills": skills
    }
